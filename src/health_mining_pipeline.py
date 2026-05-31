import argparse
import json
import os
import sqlite3
from dataclasses import dataclass
from itertools import combinations
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Set, Tuple

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from matplotlib import font_manager
from matplotlib.backends.backend_pdf import PdfPages
from sklearn.compose import ColumnTransformer
from sklearn.impute import SimpleImputer
from sklearn.metrics import (
    accuracy_score,
    classification_report,
    confusion_matrix,
    f1_score,
    precision_score,
    recall_score,
)
from sklearn.model_selection import train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.tree import DecisionTreeClassifier, export_text


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = PROJECT_ROOT / "data"
RESULTS_DIR = PROJECT_ROOT / "results"
DOCS_DIR = PROJECT_ROOT / "docs"


CONTINUOUS_COLUMNS = [
    "RIDAGEYR",
    "BMXBMI",
    "BMXWAIST",
    "BPXSY1",
    "BPXSY2",
    "BPXSY3",
    "BPXDI1",
    "BPXDI2",
    "BPXDI3",
    "LBXGH",
    "LBXTC",
    "LBDHDD",
    "LBXSTR",
    "LBXSGL",
    "DR1TKCAL",
    "DR1TSODI",
    "DR1TSUGR",
    "DR1TFIBE",
    "PAD680",
    "INDFMPIR",
]

CAT_COLUMNS = [
    "SEQN",
    "RIAGENDR",
    "RIDRETH1",
    "DMDEDUC2",
    "DIQ010",
    "BPQ020",
    "BPQ080",
    "SMQ020",
    "ALQ101",
    "PAQ605",
    "PAQ620",
    "HUQ010",
    "HSD010",
    "RXDDRUG",
    "RXDRSC1",
    "RXDRSD1",
    "RXDCOUNT",
]

YES_NO_UNKNOWN_COLUMNS = [
    "DIQ010",
    "BPQ020",
    "BPQ080",
    "SMQ020",
    "ALQ101",
    "PAQ605",
    "PAQ620",
]

TARGET_OUTCOMES = {
    "Risk=High",
    "Diabetes=Yes",
    "Hypertension=Yes",
    "Cholesterol=High",
    "Rx=Insulin",
    "Rx=Statin",
    "Rx=Antihypertensive",
}


@dataclass
class PipelineArtifacts:
    output_dir: Path
    warehouse_path: Path
    processed_path: Path
    rules_path: Path
    classification_metrics_path: Path
    report_html_path: Path
    report_md_path: Path
    report_docx_path: Optional[Path]
    report_pdf_path: Optional[Path]
    summary: Dict[str, object]


def ensure_dirs() -> None:
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    DOCS_DIR.mkdir(parents=True, exist_ok=True)


def _existing_columns(columns: Sequence[str], frame: pd.DataFrame) -> List[str]:
    return [column for column in columns if column in frame.columns]


def load_raw_data(data_path: Optional[Path] = None, nrows: Optional[int] = None) -> pd.DataFrame:
    """Load the merged NHANES dataset used by the system."""
    path = Path(data_path) if data_path else DATA_DIR / "nhanes_processed.csv"
    if not path.exists():
        raise FileNotFoundError(f"Dataset not found: {path}")
    return pd.read_csv(path, nrows=nrows, low_memory=False)


def clean_unknown_codes(df: pd.DataFrame) -> pd.DataFrame:
    cleaned = df.copy()
    for column in YES_NO_UNKNOWN_COLUMNS:
        if column in cleaned.columns:
            cleaned[column] = pd.to_numeric(cleaned[column], errors="coerce")
            cleaned.loc[cleaned[column].isin([7, 9]), column] = np.nan

    for column in ["DMDEDUC2", "HUQ010", "HSD010"]:
        if column in cleaned.columns:
            cleaned[column] = pd.to_numeric(cleaned[column], errors="coerce")
            cleaned.loc[cleaned[column].isin([7, 9]), column] = np.nan

    if "PAD680" in cleaned.columns:
        cleaned["PAD680"] = pd.to_numeric(cleaned["PAD680"], errors="coerce")
        cleaned.loc[cleaned["PAD680"].isin([7777, 9999]), "PAD680"] = np.nan
    return cleaned


def aggregate_subject_level(raw_df: pd.DataFrame) -> pd.DataFrame:
    """Collapse medication-level duplicate rows into one participant-level record."""
    if "SEQN" not in raw_df.columns:
        raise ValueError("The dataset must contain SEQN, the NHANES participant identifier.")

    selected = _existing_columns(CONTINUOUS_COLUMNS + CAT_COLUMNS, raw_df)
    df = raw_df[selected].copy()
    df = clean_unknown_codes(df)

    for column in _existing_columns(CONTINUOUS_COLUMNS + YES_NO_UNKNOWN_COLUMNS, df):
        df[column] = pd.to_numeric(df[column], errors="coerce")

    first_columns = [
        column
        for column in df.columns
        if column not in {"SEQN", "RXDDRUG", "RXDRSC1", "RXDRSD1", "RXDCOUNT"}
    ]
    subject = df.groupby("SEQN", as_index=False)[first_columns].first()

    if "RXDDRUG" in df.columns:
        meds = (
            df.groupby("SEQN")["RXDDRUG"]
            .apply(lambda values: "|".join(_unique_clean_strings(values)))
            .rename("medications")
            .reset_index()
        )
        subject = subject.merge(meds, on="SEQN", how="left")
    else:
        subject["medications"] = ""

    if "RXDRSD1" in df.columns:
        indications = (
            df.groupby("SEQN")["RXDRSD1"]
            .apply(lambda values: "|".join(_unique_clean_strings(values)))
            .rename("medication_indications")
            .reset_index()
        )
        subject = subject.merge(indications, on="SEQN", how="left")
    else:
        subject["medication_indications"] = ""

    if "RXDCOUNT" in df.columns:
        drug_count = (
            pd.to_numeric(df["RXDCOUNT"], errors="coerce")
            .groupby(df["SEQN"])
            .max()
            .rename("drug_count")
            .reset_index()
        )
        subject = subject.merge(drug_count, on="SEQN", how="left")
    else:
        subject["drug_count"] = 0

    return subject


def _unique_clean_strings(values: Iterable[object]) -> List[str]:
    invalid = {"", "99999", "77777", "REFUSED", "DON'T KNOW", "DONT KNOW", "NAN"}
    result: List[str] = []
    seen: Set[str] = set()
    for value in values:
        if pd.isna(value):
            continue
        text = str(value).strip()
        upper = text.upper()
        if upper in invalid:
            continue
        if upper not in seen:
            result.append(text)
            seen.add(upper)
    return result


def clip_continuous_outliers(df: pd.DataFrame) -> pd.DataFrame:
    clipped = df.copy()
    plausible_ranges = {
        "RIDAGEYR": (0, 120),
        "BMXBMI": (10, 90),
        "BMXWAIST": (35, 200),
        "BPXSY1": (60, 260),
        "BPXSY2": (60, 260),
        "BPXSY3": (60, 260),
        "BPXDI1": (30, 160),
        "BPXDI2": (30, 160),
        "BPXDI3": (30, 160),
        "LBXGH": (3, 20),
        "LBXTC": (50, 900),
        "LBDHDD": (5, 200),
        "LBXSTR": (20, 2000),
        "LBXSGL": (40, 700),
        "DR1TKCAL": (300, 9000),
        "DR1TSODI": (0, 25000),
        "DR1TSUGR": (0, 1500),
        "DR1TFIBE": (0, 160),
        "PAD680": (0, 1440),
        "INDFMPIR": (0, 5),
    }
    for column, (lower, upper) in plausible_ranges.items():
        if column in clipped.columns:
            clipped[column] = pd.to_numeric(clipped[column], errors="coerce")
            clipped.loc[(clipped[column] < lower) | (clipped[column] > upper), column] = np.nan
    return clipped


def add_health_features(subject_df: pd.DataFrame, adult_only: bool = True) -> pd.DataFrame:
    df = clip_continuous_outliers(subject_df)
    if adult_only and "RIDAGEYR" in df.columns:
        df = df[df["RIDAGEYR"] >= 18].copy()

    df["gender"] = df.get("RIAGENDR", pd.Series(index=df.index)).map({1: "Male", 2: "Female"}).fillna("Unknown")
    df["age"] = pd.to_numeric(df.get("RIDAGEYR"), errors="coerce")
    df["age_group"] = pd.cut(
        df["age"],
        bins=[17, 29, 44, 59, 120],
        labels=["18-29", "30-44", "45-59", "60+"],
        include_lowest=True,
    ).astype("object")

    df["bmi"] = pd.to_numeric(df.get("BMXBMI"), errors="coerce")
    df["bmi_group"] = pd.cut(
        df["bmi"],
        bins=[0, 18.5, 25, 30, 200],
        labels=["Underweight", "Normal", "Overweight", "Obese"],
        include_lowest=True,
    ).astype("object")

    systolic_cols = _existing_columns(["BPXSY1", "BPXSY2", "BPXSY3"], df)
    diastolic_cols = _existing_columns(["BPXDI1", "BPXDI2", "BPXDI3"], df)
    df["systolic_bp"] = df[systolic_cols].mean(axis=1, skipna=True) if systolic_cols else np.nan
    df["diastolic_bp"] = df[diastolic_cols].mean(axis=1, skipna=True) if diastolic_cols else np.nan

    df["hba1c"] = pd.to_numeric(df.get("LBXGH"), errors="coerce")
    df["total_cholesterol"] = pd.to_numeric(df.get("LBXTC"), errors="coerce")
    df["hdl_cholesterol"] = pd.to_numeric(df.get("LBDHDD"), errors="coerce")
    df["triglycerides"] = pd.to_numeric(df.get("LBXSTR"), errors="coerce")
    df["glucose"] = pd.to_numeric(df.get("LBXSGL"), errors="coerce")
    df["sodium_mg"] = pd.to_numeric(df.get("DR1TSODI"), errors="coerce")
    df["sugar_g"] = pd.to_numeric(df.get("DR1TSUGR"), errors="coerce")
    df["fiber_g"] = pd.to_numeric(df.get("DR1TFIBE"), errors="coerce")
    df["calories"] = pd.to_numeric(df.get("DR1TKCAL"), errors="coerce")
    df["sedentary_minutes"] = pd.to_numeric(df.get("PAD680"), errors="coerce")
    df["income_poverty_ratio"] = pd.to_numeric(df.get("INDFMPIR"), errors="coerce")

    df["diabetes_self_report"] = (pd.to_numeric(df.get("DIQ010"), errors="coerce") == 1).astype(int)
    df["hypertension_self_report"] = (pd.to_numeric(df.get("BPQ020"), errors="coerce") == 1).astype(int)
    df["cholesterol_self_report"] = (pd.to_numeric(df.get("BPQ080"), errors="coerce") == 1).astype(int)
    df["ever_smoker"] = (pd.to_numeric(df.get("SMQ020"), errors="coerce") == 1).astype(int)
    df["regular_alcohol"] = (pd.to_numeric(df.get("ALQ101"), errors="coerce") == 1).astype(int)
    df["vigorous_activity"] = (pd.to_numeric(df.get("PAQ605"), errors="coerce") == 1).astype(int)
    df["moderate_activity"] = (pd.to_numeric(df.get("PAQ620"), errors="coerce") == 1).astype(int)

    df["waist_high"] = np.where(
        df["gender"].eq("Male"),
        df.get("BMXWAIST", np.nan) > 102,
        np.where(df["gender"].eq("Female"), df.get("BMXWAIST", np.nan) > 88, False),
    ).astype(int)
    df["obesity"] = (df["bmi"] >= 30).astype(int)
    df["hypertension"] = (
        (df["hypertension_self_report"] == 1)
        | (df["systolic_bp"] >= 130)
        | (df["diastolic_bp"] >= 80)
    ).astype(int)
    df["diabetes"] = (
        (df["diabetes_self_report"] == 1)
        | (df["hba1c"] >= 6.5)
        | (df["glucose"] >= 126)
    ).astype(int)
    df["prediabetes"] = ((df["hba1c"] >= 5.7) & (df["hba1c"] < 6.5)).astype(int)
    df["high_cholesterol"] = (
        (df["cholesterol_self_report"] == 1)
        | (df["total_cholesterol"] >= 240)
        | ((df["gender"].eq("Male")) & (df["hdl_cholesterol"] < 40))
        | ((df["gender"].eq("Female")) & (df["hdl_cholesterol"] < 50))
        | (df["triglycerides"] >= 150)
    ).astype(int)
    df["high_sodium"] = (df["sodium_mg"] > 2300).astype(int)
    df["high_sugar"] = (df["sugar_g"] > 75).astype(int)
    df["low_fiber"] = (df["fiber_g"] < 14).astype(int)
    df["sedentary_high"] = (df["sedentary_minutes"] >= 600).astype(int)
    df["older_adult"] = (df["age"] >= 60).astype(int)

    df["rx_insulin"] = df["medications"].fillna("").str.contains("INSULIN", case=False, regex=False).astype(int)
    df["rx_metformin"] = df["medications"].fillna("").str.contains("METFORMIN", case=False, regex=False).astype(int)
    df["rx_statin"] = df["medications"].fillna("").str.contains(
        "SIMVASTATIN|ATORVASTATIN|ROSUVASTATIN|PRAVASTATIN|LOVASTATIN",
        case=False,
        regex=True,
    ).astype(int)
    df["rx_antihypertensive"] = df["medications"].fillna("").str.contains(
        "LISINOPRIL|LOSARTAN|OLMESARTAN|VALSARTAN|AMLODIPINE|HYDROCHLOROTHIAZIDE|"
        "ATENOLOL|METOPROLOL|CARVEDILOL|FUROSEMIDE",
        case=False,
        regex=True,
    ).astype(int)

    risk_factors = [
        "obesity",
        "waist_high",
        "hypertension",
        "diabetes",
        "prediabetes",
        "high_cholesterol",
        "ever_smoker",
        "high_sodium",
        "low_fiber",
        "sedentary_high",
        "older_adult",
    ]
    df["risk_score"] = df[risk_factors].sum(axis=1)
    df["risk_level"] = pd.cut(
        df["risk_score"],
        bins=[-1, 2, 4, 20],
        labels=["Low", "Medium", "High"],
    ).astype("object")
    df["high_cardiometabolic_risk"] = (df["risk_level"] == "High").astype(int)

    output_columns = [
        "SEQN",
        "age",
        "gender",
        "age_group",
        "bmi",
        "bmi_group",
        "systolic_bp",
        "diastolic_bp",
        "hba1c",
        "glucose",
        "total_cholesterol",
        "hdl_cholesterol",
        "triglycerides",
        "calories",
        "sodium_mg",
        "sugar_g",
        "fiber_g",
        "sedentary_minutes",
        "income_poverty_ratio",
        "diabetes_self_report",
        "hypertension_self_report",
        "cholesterol_self_report",
        "ever_smoker",
        "regular_alcohol",
        "vigorous_activity",
        "moderate_activity",
        "waist_high",
        "obesity",
        "hypertension",
        "diabetes",
        "prediabetes",
        "high_cholesterol",
        "high_sodium",
        "high_sugar",
        "low_fiber",
        "sedentary_high",
        "older_adult",
        "rx_insulin",
        "rx_metformin",
        "rx_statin",
        "rx_antihypertensive",
        "drug_count",
        "medications",
        "medication_indications",
        "risk_score",
        "risk_level",
        "high_cardiometabolic_risk",
    ]
    return df[_existing_columns(output_columns, df)].reset_index(drop=True)


def build_aggregation_tables(processed: pd.DataFrame) -> Dict[str, pd.DataFrame]:
    aggregations: Dict[str, pd.DataFrame] = {}
    group_specs = {
        "age_group": "age_group",
        "gender": "gender",
        "bmi_group": "bmi_group",
        "risk_level": "risk_level",
    }
    for name, column in group_specs.items():
        table = (
            processed.groupby(column, dropna=False)
            .agg(
                participants=("SEQN", "count"),
                mean_age=("age", "mean"),
                mean_bmi=("bmi", "mean"),
                mean_systolic_bp=("systolic_bp", "mean"),
                mean_hba1c=("hba1c", "mean"),
                mean_total_cholesterol=("total_cholesterol", "mean"),
                high_risk_rate=("high_cardiometabolic_risk", "mean"),
                diabetes_rate=("diabetes", "mean"),
                hypertension_rate=("hypertension", "mean"),
                high_cholesterol_rate=("high_cholesterol", "mean"),
                obesity_rate=("obesity", "mean"),
            )
            .reset_index()
        )
        aggregations[name] = table
    return aggregations


def transaction_from_row(row: pd.Series) -> Set[str]:
    items: Set[str] = set()
    if pd.notna(row.get("age_group")):
        items.add(f"Age={row['age_group']}")
    if pd.notna(row.get("gender")) and row.get("gender") != "Unknown":
        items.add(f"Sex={row['gender']}")
    if pd.notna(row.get("bmi_group")):
        items.add(f"BMI={row['bmi_group']}")
    if row.get("hypertension", 0) == 1:
        items.add("Hypertension=Yes")
    else:
        items.add("Hypertension=No")
    if row.get("diabetes", 0) == 1:
        items.add("Diabetes=Yes")
    elif row.get("prediabetes", 0) == 1:
        items.add("Diabetes=Prediabetes")
    else:
        items.add("Diabetes=No")
    if row.get("high_cholesterol", 0) == 1:
        items.add("Cholesterol=High")
    else:
        items.add("Cholesterol=Normal")
    if row.get("ever_smoker", 0) == 1:
        items.add("Smoking=Ever")
    else:
        items.add("Smoking=Never")
    if row.get("regular_alcohol", 0) == 1:
        items.add("Alcohol=Regular")
    if row.get("high_sodium", 0) == 1:
        items.add("Diet=SodiumHigh")
    if row.get("high_sugar", 0) == 1:
        items.add("Diet=SugarHigh")
    if row.get("low_fiber", 0) == 1:
        items.add("Diet=FiberLow")
    if row.get("sedentary_high", 0) == 1:
        items.add("Activity=SedentaryHigh")
    if row.get("vigorous_activity", 0) == 1 or row.get("moderate_activity", 0) == 1:
        items.add("Activity=Active")
    if row.get("rx_insulin", 0) == 1:
        items.add("Rx=Insulin")
    if row.get("rx_metformin", 0) == 1:
        items.add("Rx=Metformin")
    if row.get("rx_statin", 0) == 1:
        items.add("Rx=Statin")
    if row.get("rx_antihypertensive", 0) == 1:
        items.add("Rx=Antihypertensive")
    items.add("Risk=High" if row.get("high_cardiometabolic_risk", 0) == 1 else "Risk=NotHigh")
    return items


def apriori_rules(
    transactions: Sequence[Set[str]],
    min_support: float = 0.08,
    min_confidence: float = 0.55,
    max_len: int = 3,
    target_outcomes: Optional[Set[str]] = None,
) -> Tuple[pd.DataFrame, pd.DataFrame]:
    if not transactions:
        return pd.DataFrame(), pd.DataFrame()

    target_outcomes = target_outcomes or TARGET_OUTCOMES
    n_transactions = len(transactions)
    support_map: Dict[frozenset, float] = {}

    item_counts: Dict[frozenset, int] = {}
    for transaction in transactions:
        for item in transaction:
            key = frozenset([item])
            item_counts[key] = item_counts.get(key, 0) + 1

    current_level = {
        itemset
        for itemset, count in item_counts.items()
        if count / n_transactions >= min_support
    }
    for itemset in current_level:
        support_map[itemset] = item_counts[itemset] / n_transactions

    all_frequent = set(current_level)
    k = 2
    while current_level and k <= max_len:
        candidates: Set[frozenset] = set()
        current_list = sorted(current_level, key=lambda itemset: sorted(itemset))
        for left, right in combinations(current_list, 2):
            union = left | right
            if len(union) != k:
                continue
            if all(frozenset(subset) in all_frequent for subset in combinations(union, k - 1)):
                candidates.add(union)

        candidate_counts = {candidate: 0 for candidate in candidates}
        for transaction in transactions:
            for candidate in candidates:
                if candidate.issubset(transaction):
                    candidate_counts[candidate] += 1

        current_level = {
            candidate
            for candidate, count in candidate_counts.items()
            if count / n_transactions >= min_support
        }
        for itemset in current_level:
            support_map[itemset] = candidate_counts[itemset] / n_transactions
        all_frequent.update(current_level)
        k += 1

    itemset_rows = [
        {"items": " + ".join(sorted(itemset)), "item_count": len(itemset), "support": support}
        for itemset, support in support_map.items()
    ]
    itemsets_df = pd.DataFrame(itemset_rows).sort_values(
        ["item_count", "support", "items"], ascending=[True, False, True]
    )

    rule_rows = []
    for itemset, support in support_map.items():
        if len(itemset) < 2:
            continue
        for size in range(1, len(itemset)):
            for antecedent_tuple in combinations(itemset, size):
                antecedent = frozenset(antecedent_tuple)
                consequent = itemset - antecedent
                if len(consequent) != 1:
                    continue
                consequent_item = next(iter(consequent))
                if consequent_item not in target_outcomes:
                    continue
                antecedent_support = support_map.get(antecedent)
                consequent_support = support_map.get(frozenset(consequent))
                if not antecedent_support or not consequent_support:
                    continue
                confidence = support / antecedent_support
                if confidence < min_confidence:
                    continue
                lift = confidence / consequent_support
                rule_rows.append(
                    {
                        "antecedent": " + ".join(sorted(antecedent)),
                        "consequent": consequent_item,
                        "support": support,
                        "confidence": confidence,
                        "lift": lift,
                        "antecedent_support": antecedent_support,
                        "consequent_support": consequent_support,
                    }
                )

    rules_df = pd.DataFrame(rule_rows)
    if not rules_df.empty:
        rules_df = rules_df.sort_values(["lift", "confidence", "support"], ascending=False).reset_index(drop=True)
    return itemsets_df.reset_index(drop=True), rules_df


def run_classification(
    processed: pd.DataFrame,
    max_depth: int = 5,
    random_state: int = 42,
) -> Tuple[Dict[str, float], pd.DataFrame, pd.DataFrame, str]:
    numeric_features = [
        "age",
        "bmi",
        "systolic_bp",
        "diastolic_bp",
        "hba1c",
        "glucose",
        "total_cholesterol",
        "hdl_cholesterol",
        "triglycerides",
        "sodium_mg",
        "sugar_g",
        "fiber_g",
        "sedentary_minutes",
        "income_poverty_ratio",
        "drug_count",
    ]
    categorical_features = ["gender", "age_group", "bmi_group"]
    binary_features = [
        "ever_smoker",
        "regular_alcohol",
        "vigorous_activity",
        "moderate_activity",
        "rx_insulin",
        "rx_metformin",
        "rx_statin",
        "rx_antihypertensive",
    ]

    numeric_features = _existing_columns(numeric_features, processed)
    categorical_features = _existing_columns(categorical_features, processed)
    binary_features = _existing_columns(binary_features, processed)
    feature_columns = numeric_features + categorical_features + binary_features
    model_df = processed[feature_columns + ["high_cardiometabolic_risk"]].dropna(
        subset=["high_cardiometabolic_risk"]
    )
    X = model_df[feature_columns]
    y = model_df["high_cardiometabolic_risk"].astype(int)

    preprocessor = ColumnTransformer(
        transformers=[
            ("num", Pipeline([("imputer", SimpleImputer(strategy="median")), ("scaler", StandardScaler())]), numeric_features),
            ("cat", Pipeline([("imputer", SimpleImputer(strategy="most_frequent")), ("onehot", OneHotEncoder(handle_unknown="ignore"))]), categorical_features),
            ("bin", Pipeline([("imputer", SimpleImputer(strategy="most_frequent"))]), binary_features),
        ]
    )
    classifier = DecisionTreeClassifier(
        criterion="entropy",
        max_depth=max_depth,
        min_samples_leaf=40,
        class_weight="balanced",
        random_state=random_state,
    )
    model = Pipeline([("preprocess", preprocessor), ("classifier", classifier)])

    stratify = y if y.nunique() > 1 else None
    X_train, X_test, y_train, y_test = train_test_split(
        X,
        y,
        test_size=0.3,
        random_state=random_state,
        stratify=stratify,
    )
    model.fit(X_train, y_train)
    predictions = model.predict(X_test)

    metrics = {
        "samples": float(len(model_df)),
        "train_samples": float(len(X_train)),
        "test_samples": float(len(X_test)),
        "positive_rate": float(y.mean()),
        "accuracy": float(accuracy_score(y_test, predictions)),
        "precision": float(precision_score(y_test, predictions, zero_division=0)),
        "recall": float(recall_score(y_test, predictions, zero_division=0)),
        "f1": float(f1_score(y_test, predictions, zero_division=0)),
    }

    cm = confusion_matrix(y_test, predictions, labels=[0, 1])
    confusion_df = pd.DataFrame(
        cm,
        index=["Actual_NotHigh", "Actual_High"],
        columns=["Pred_NotHigh", "Pred_High"],
    )

    feature_names = get_feature_names(model, numeric_features, categorical_features, binary_features)
    importances = model.named_steps["classifier"].feature_importances_
    importance_df = (
        pd.DataFrame({"feature": feature_names, "importance": importances})
        .sort_values("importance", ascending=False)
        .reset_index(drop=True)
    )

    tree_text = export_text(
        model.named_steps["classifier"],
        feature_names=list(feature_names),
        max_depth=max_depth,
    )
    report = classification_report(y_test, predictions, zero_division=0)
    tree_text = f"{report}\n\nDecision tree rules:\n{tree_text}"
    return metrics, confusion_df, importance_df, tree_text


def get_feature_names(
    model: Pipeline,
    numeric_features: List[str],
    categorical_features: List[str],
    binary_features: List[str],
) -> np.ndarray:
    names: List[str] = []
    names.extend(numeric_features)
    if categorical_features:
        encoder = model.named_steps["preprocess"].named_transformers_["cat"].named_steps["onehot"]
        names.extend(list(encoder.get_feature_names_out(categorical_features)))
    names.extend(binary_features)
    return np.array(names)


def save_csv_outputs(
    processed: pd.DataFrame,
    aggregations: Dict[str, pd.DataFrame],
    itemsets: pd.DataFrame,
    rules: pd.DataFrame,
    metrics: Dict[str, float],
    confusion: pd.DataFrame,
    importances: pd.DataFrame,
    output_dir: Path,
) -> None:
    processed.to_csv(output_dir / "nhanes_health_subjects.csv", index=False, encoding="utf-8-sig")
    for name, table in aggregations.items():
        table.to_csv(output_dir / f"aggregation_by_{name}.csv", index=False, encoding="utf-8-sig")
    itemsets.to_csv(output_dir / "association_frequent_itemsets.csv", index=False, encoding="utf-8-sig")
    rules.to_csv(output_dir / "association_rules.csv", index=False, encoding="utf-8-sig")
    pd.DataFrame([metrics]).to_csv(output_dir / "classification_metrics.csv", index=False, encoding="utf-8-sig")
    confusion.to_csv(output_dir / "classification_confusion_matrix.csv", encoding="utf-8-sig")
    importances.to_csv(output_dir / "classification_feature_importance.csv", index=False, encoding="utf-8-sig")


def save_visualizations(
    aggregations: Dict[str, pd.DataFrame],
    confusion: pd.DataFrame,
    importances: pd.DataFrame,
    output_dir: Path,
) -> None:
    sns.set_theme(style="whitegrid", font="DejaVu Sans")

    age_table = aggregations.get("age_group")
    if age_table is not None and not age_table.empty:
        plt.figure(figsize=(8, 5))
        plot_df = age_table.copy()
        plot_df["high_risk_rate_pct"] = plot_df["high_risk_rate"] * 100
        sns.barplot(data=plot_df, x="age_group", y="high_risk_rate_pct", color="#2563eb")
        plt.title("High Cardiometabolic Risk Rate by Age Group")
        plt.xlabel("Age group")
        plt.ylabel("High risk rate (%)")
        plt.tight_layout()
        plt.savefig(output_dir / "aggregation_risk_by_age.png", dpi=180)
        plt.close()

        disease = plot_df[["age_group", "diabetes_rate", "hypertension_rate", "high_cholesterol_rate"]].melt(
            id_vars="age_group",
            var_name="indicator",
            value_name="rate",
        )
        disease["rate"] = disease["rate"] * 100
        plt.figure(figsize=(9, 5))
        sns.lineplot(data=disease, x="age_group", y="rate", hue="indicator", marker="o")
        plt.title("Disease Indicator Rates by Age Group")
        plt.xlabel("Age group")
        plt.ylabel("Rate (%)")
        plt.tight_layout()
        plt.savefig(output_dir / "aggregation_disease_rates_by_age.png", dpi=180)
        plt.close()

    bmi_table = aggregations.get("bmi_group")
    if bmi_table is not None and not bmi_table.empty:
        plt.figure(figsize=(8, 5))
        plot_df = bmi_table.copy()
        plot_df["high_risk_rate_pct"] = plot_df["high_risk_rate"] * 100
        sns.barplot(data=plot_df, x="bmi_group", y="high_risk_rate_pct", color="#059669")
        plt.title("High Cardiometabolic Risk Rate by BMI Group")
        plt.xlabel("BMI group")
        plt.ylabel("High risk rate (%)")
        plt.tight_layout()
        plt.savefig(output_dir / "aggregation_risk_by_bmi.png", dpi=180)
        plt.close()

    plt.figure(figsize=(5, 4))
    sns.heatmap(confusion, annot=True, fmt="d", cmap="Blues", cbar=False)
    plt.title("Decision Tree Confusion Matrix")
    plt.tight_layout()
    plt.savefig(output_dir / "classification_confusion_matrix.png", dpi=180)
    plt.close()

    top_importances = importances.head(12)
    plt.figure(figsize=(9, 6))
    sns.barplot(data=top_importances, y="feature", x="importance", color="#7c3aed")
    plt.title("Top Decision Tree Feature Importances")
    plt.xlabel("Importance")
    plt.ylabel("")
    plt.tight_layout()
    plt.savefig(output_dir / "classification_feature_importance.png", dpi=180)
    plt.close()


def write_sqlite_warehouse(
    raw_subject: pd.DataFrame,
    processed: pd.DataFrame,
    aggregations: Dict[str, pd.DataFrame],
    itemsets: pd.DataFrame,
    rules: pd.DataFrame,
    metrics: Dict[str, float],
    confusion: pd.DataFrame,
    importances: pd.DataFrame,
    output_dir: Path,
) -> Path:
    warehouse_path = output_dir / "nhanes_health_warehouse.db"
    if warehouse_path.exists():
        warehouse_path.unlink()
    with sqlite3.connect(warehouse_path) as conn:
        raw_subject.to_sql("ods_nhanes_subject_raw", conn, index=False)
        processed.to_sql("dwd_health_subjects", conn, index=False)
        for name, table in aggregations.items():
            table.to_sql(f"ads_aggregation_by_{name}", conn, index=False)
        itemsets.to_sql("dm_association_frequent_itemsets", conn, index=False)
        rules.to_sql("dm_association_rules", conn, index=False)
        pd.DataFrame([metrics]).to_sql("dm_classification_metrics", conn, index=False)
        confusion.reset_index(names="actual").to_sql("dm_classification_confusion_matrix", conn, index=False)
        importances.to_sql("dm_classification_feature_importance", conn, index=False)
    return warehouse_path


def generate_warehouse_design(processed: pd.DataFrame, output_dir: Path) -> Path:
    doc_path = DOCS_DIR / "data_warehouse_design.md"
    fields = [
        ("ODS", "ods_nhanes_subject_raw", "按 SEQN 汇总后的原始受试者记录，保留人口学、体检、实验室、问卷、用药字段。"),
        ("DWD", "dwd_health_subjects", "清洗、规约和医学特征工程后的受试者宽表。"),
        ("ADS", "ads_aggregation_by_age_group", "按年龄分组的健康风险聚合指标。"),
        ("ADS", "ads_aggregation_by_gender", "按性别分组的健康风险聚合指标。"),
        ("ADS", "ads_aggregation_by_bmi_group", "按 BMI 分层的健康风险聚合指标。"),
        ("DM", "dm_association_rules", "Apriori 关联规则结果。"),
        ("DM", "dm_classification_metrics", "决策树分类评价指标。"),
        ("DM", "dm_classification_feature_importance", "决策树特征重要性。"),
    ]
    lines = [
        "# NHANES 心代谢风险数据仓库设计",
        "",
        "## 分层设计",
        "",
        "| 层级 | 表名 | 说明 |",
        "| --- | --- | --- |",
    ]
    lines.extend([f"| {layer} | `{table}` | {desc} |" for layer, table, desc in fields])
    lines.extend(
        [
            "",
            "## DWD 核心字段",
            "",
            "| 字段 | 含义 |",
            "| --- | --- |",
        ]
    )
    descriptions = {
        "SEQN": "NHANES 脱敏受试者编号",
        "age": "年龄",
        "gender": "性别",
        "age_group": "年龄分组",
        "bmi": "体质指数",
        "bmi_group": "BMI 分层",
        "systolic_bp": "平均收缩压",
        "diastolic_bp": "平均舒张压",
        "hba1c": "糖化血红蛋白",
        "total_cholesterol": "总胆固醇",
        "hdl_cholesterol": "HDL 胆固醇",
        "hypertension": "是否高血压风险",
        "diabetes": "是否糖尿病风险",
        "high_cholesterol": "是否高胆固醇风险",
        "risk_score": "心代谢风险因素计数",
        "risk_level": "心代谢风险分层",
        "high_cardiometabolic_risk": "分类目标：是否高心代谢风险",
    }
    for column in processed.columns:
        lines.append(f"| `{column}` | {descriptions.get(column, '清洗或派生后的分析字段')} |")
    lines.extend(
        [
            "",
            "## 数据流",
            "",
            "1. ODS：读取 `data/nhanes_processed.csv`，按 `SEQN` 汇总重复的用药记录。",
            "2. DWD：清洗未知编码、异常范围和缺失值，生成 BMI、血压、糖尿病、胆固醇、饮食与用药特征。",
            "3. ADS：按年龄、性别、BMI 和风险等级进行聚合分析。",
            "4. DM：执行 Apriori 关联规则和决策树分类，结果写回 SQLite 与 CSV。",
        ]
    )
    doc_path.write_text("\n".join(lines), encoding="utf-8")
    return doc_path


def generate_course_report(
    processed: pd.DataFrame,
    aggregations: Dict[str, pd.DataFrame],
    rules: pd.DataFrame,
    metrics: Dict[str, float],
    importances: pd.DataFrame,
    output_dir: Path,
) -> Tuple[Path, Optional[Path]]:
    md_path = DOCS_DIR / "course_report.md"
    top_rules = rules.head(10).copy()
    top_importances = importances.head(10).copy()
    age_agg = aggregations.get("age_group", pd.DataFrame()).copy()
    bmi_agg = aggregations.get("bmi_group", pd.DataFrame()).copy()

    lines = [
        "# 基于 NHANES 数据的心代谢风险分层、关联规则与分类预测系统",
        "",
        "## 摘要",
        "",
        (
            "本项目围绕医学健康管理场景，使用 NHANES 公开脱敏数据，构建心代谢风险数据仓库，"
            "实现数据预处理、聚合分析、Apriori 关联规则挖掘和决策树分类预测，并通过 Streamlit 系统展示结果。"
        ),
        "",
        "## 数据集与场景",
        "",
        "- 数据来源：CDC NHANES 公开调查数据，仓库内已提供合并后的 `data/nhanes_processed.csv`。",
        f"- 分析对象：清洗后成人受试者 {len(processed):,} 条。",
        "- 医学场景：识别具有肥胖、高血压、糖尿病、高胆固醇、不良生活方式等因素的心代谢高风险人群。",
        "",
        "## 数据预处理与数据仓库",
        "",
        "- 按 `SEQN` 将多行用药记录规约为受试者级宽表。",
        "- 将 7、9、7777、9999 等 NHANES 未知或拒答编码转为空值。",
        "- 按医学合理范围剔除异常值，并生成年龄组、BMI 分层、血压、糖化血红蛋白、胆固醇、饮食、活动和用药标签。",
        "- SQLite 数据仓库路径：`results/nhanes_health_warehouse.db`。",
        "",
        "## 聚合分析",
        "",
        _markdown_table(age_agg.head(10)),
        "",
        _markdown_table(bmi_agg.head(10)),
        "",
        "## 关联规则算法",
        "",
        f"- Apriori 默认参数：最小支持度 0.08，最小置信度 0.55，最大项集长度 3。",
        "- 规则重点筛选高风险、糖尿病、高血压、高胆固醇和典型用药相关后件。",
        "",
        _markdown_table(top_rules),
        "",
        "## 分类算法",
        "",
        "- 算法：决策树（信息增益/熵准则），训练集与测试集比例 7:3。",
        f"- Accuracy：{metrics.get('accuracy', 0):.3f}",
        f"- Precision：{metrics.get('precision', 0):.3f}",
        f"- Recall：{metrics.get('recall', 0):.3f}",
        f"- F1：{metrics.get('f1', 0):.3f}",
        "",
        "### 重要特征",
        "",
        _markdown_table(top_importances),
        "",
        "## 系统开发",
        "",
        "- 命令行入口：`python src/main.py`。",
        "- Streamlit 入口：`streamlit run src/app.py`。",
        "- 系统模块：数据导入、数据预处理、数据仓库、聚合分析、关联规则、分类预测、结果导出。",
        "",
        "## 测试用例",
        "",
        "| 测试项 | 操作 | 预期结果 |",
        "| --- | --- | --- |",
        "| 数据导入 | 运行默认数据或上传 CSV | 展示样本量和字段数 |",
        "| 预处理 | 点击运行分析 | 生成受试者级清洗数据 |",
        "| 数据仓库 | 查看结果文件 | SQLite 数据库包含 ODS/DWD/ADS/DM 表 |",
        "| 聚合分析 | 查看聚合页面 | 年龄、性别、BMI 分组指标可视化 |",
        "| 关联规则 | 调整支持度/置信度并运行 | 生成规则表，包含支持度、置信度、提升度 |",
        "| 分类预测 | 查看分类页面 | 输出准确率、召回率、F1 和混淆矩阵 |",
        "",
        "## 总结与展望",
        "",
        "项目完成了医学数据从导入、清洗、仓库分层、算法挖掘到系统展示的闭环。后续可进一步引入抽样权重、更多 NHANES 周期数据和临床验证规则，提高医学解释性。",
        "",
        "## 参考文献",
        "",
        "- CDC National Health and Nutrition Examination Survey (NHANES).",
        "- Han, J., Kamber, M., Pei, J. Data Mining: Concepts and Techniques.",
        "- scikit-learn documentation: DecisionTreeClassifier.",
        "- Streamlit documentation.",
    ]
    md_path.write_text("\n".join(lines), encoding="utf-8")

    docx_path = None
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("基于 NHANES 数据的心代谢风险分层、关联规则与分类预测系统", level=0)
        for line in lines[2:]:
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=2)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("|") or not line.strip():
                continue
            else:
                doc.add_paragraph(line)
        docx_path = DOCS_DIR / "course_report.docx"
        doc.save(docx_path)
    except Exception:
        docx_path = None

    return md_path, docx_path


def _get_pdf_font():
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return font_manager.FontProperties(fname=str(candidate))
    return font_manager.FontProperties(family="DejaVu Sans")


def _wrap_text(text: str, width: int = 54) -> List[str]:
    lines: List[str] = []
    for paragraph in text.splitlines():
        if not paragraph:
            lines.append("")
            continue
        current = ""
        for char in paragraph:
            current += char
            if len(current) >= width:
                lines.append(current)
                current = ""
        if current:
            lines.append(current)
    return lines


def _pdf_text_page(pdf: PdfPages, title: str, body: Sequence[str], font_prop) -> None:
    fig = plt.figure(figsize=(8.27, 11.69))
    fig.patch.set_facecolor("white")
    fig.text(0.08, 0.94, title, fontproperties=font_prop, fontsize=17, weight="bold")
    y = 0.88
    for paragraph in body:
        for line in _wrap_text(paragraph):
            fig.text(0.08, y, line, fontproperties=font_prop, fontsize=10)
            y -= 0.026
            if y < 0.08:
                pdf.savefig(fig, bbox_inches="tight")
                plt.close(fig)
                fig = plt.figure(figsize=(8.27, 11.69))
                fig.patch.set_facecolor("white")
                y = 0.94
        y -= 0.012
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


def _pdf_table_page(pdf: PdfPages, title: str, df: pd.DataFrame, font_prop, fontsize: int = 8) -> None:
    fig, ax = plt.subplots(figsize=(11.69, 8.27))
    ax.axis("off")
    ax.set_title(title, fontproperties=font_prop, fontsize=15, weight="bold", pad=16)
    display = df.copy()
    for column in display.select_dtypes(include=[np.number]).columns:
        display[column] = display[column].map(lambda value: f"{value:.3f}" if pd.notna(value) else "")
    table = ax.table(
        cellText=display.astype(str).values,
        colLabels=[str(column) for column in display.columns],
        loc="center",
        cellLoc="center",
    )
    table.auto_set_font_size(False)
    table.set_fontsize(fontsize)
    table.scale(1, 1.45)
    for cell in table.get_celld().values():
        cell.get_text().set_fontproperties(font_prop)
    pdf.savefig(fig, bbox_inches="tight")
    plt.close(fig)


def generate_pdf_report(
    processed: pd.DataFrame,
    aggregations: Dict[str, pd.DataFrame],
    rules: pd.DataFrame,
    metrics: Dict[str, float],
    importances: pd.DataFrame,
    output_dir: Path,
) -> Optional[Path]:
    pdf_path = DOCS_DIR / "course_report.pdf"
    try:
        font_prop = _get_pdf_font()
        with PdfPages(pdf_path) as pdf:
            _pdf_text_page(
                pdf,
                "基于 NHANES 数据的心代谢风险数据挖掘系统",
                [
                    "摘要：本项目使用 NHANES 公开脱敏健康数据，构建心代谢风险数据仓库，完成数据预处理、聚合分析、Apriori 关联规则和决策树分类预测，并通过 Streamlit 系统展示结果。",
                    f"数据规模：全量原始合并记录经受试者级规约与成人筛选后，形成 {len(processed):,} 条成人分析记录。",
                    "医学场景：识别肥胖、高血压、糖尿病、高胆固醇、不良饮食、久坐和吸烟等因素叠加形成的心代谢高风险人群。",
                    "系统模块：数据导入、预处理与数据仓库、聚合分析、关联规则、分类预测、结果导出、系统说明。",
                    f"分类结果：Accuracy={metrics.get('accuracy', 0):.3f}，Precision={metrics.get('precision', 0):.3f}，Recall={metrics.get('recall', 0):.3f}，F1={metrics.get('f1', 0):.3f}。",
                    "数据仓库：SQLite 文件 results/nhanes_health_warehouse.db，包含 ODS、DWD、ADS、DM 四层表。",
                ],
                font_prop,
            )
            age = aggregations.get("age_group", pd.DataFrame())
            if not age.empty:
                _pdf_table_page(
                    pdf,
                    "聚合分析：按年龄分组",
                    age[
                        [
                            "age_group",
                            "participants",
                            "high_risk_rate",
                            "diabetes_rate",
                            "hypertension_rate",
                            "high_cholesterol_rate",
                        ]
                    ],
                    font_prop,
                )
            bmi = aggregations.get("bmi_group", pd.DataFrame())
            if not bmi.empty:
                _pdf_table_page(
                    pdf,
                    "聚合分析：按 BMI 分层",
                    bmi[
                        [
                            "bmi_group",
                            "participants",
                            "high_risk_rate",
                            "diabetes_rate",
                            "hypertension_rate",
                            "obesity_rate",
                        ]
                    ],
                    font_prop,
                )
            if not rules.empty:
                _pdf_table_page(
                    pdf,
                    "Apriori 关联规则 Top 10",
                    rules[["antecedent", "consequent", "support", "confidence", "lift"]].head(10),
                    font_prop,
                    fontsize=7,
                )
            if not importances.empty:
                _pdf_table_page(
                    pdf,
                    "决策树特征重要性 Top 10",
                    importances[["feature", "importance"]].head(10),
                    font_prop,
                )
        return pdf_path
    except Exception:
        return None


def _markdown_table(df: pd.DataFrame) -> str:
    if df is None or df.empty:
        return "暂无结果。"
    display = df.copy()
    for column in display.select_dtypes(include=[np.number]).columns:
        display[column] = display[column].map(lambda value: f"{value:.4f}" if pd.notna(value) else "")
    headers = [str(column) for column in display.columns]
    rows = display.astype(str).values.tolist()
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in rows)
    return "\n".join(lines)


def generate_html_report(
    processed: pd.DataFrame,
    aggregations: Dict[str, pd.DataFrame],
    rules: pd.DataFrame,
    metrics: Dict[str, float],
    importances: pd.DataFrame,
    output_dir: Path,
) -> Path:
    report_path = output_dir / "health_mining_report.html"
    top_rules_html = rules.head(12).to_html(index=False, float_format=lambda value: f"{value:.3f}") if not rules.empty else "<p>No rules generated.</p>"
    metrics_html = pd.DataFrame([metrics]).to_html(index=False, float_format=lambda value: f"{value:.3f}")
    imports_html = importances.head(12).to_html(index=False, float_format=lambda value: f"{value:.3f}")
    age_html = aggregations.get("age_group", pd.DataFrame()).to_html(index=False, float_format=lambda value: f"{value:.3f}")
    bmi_html = aggregations.get("bmi_group", pd.DataFrame()).to_html(index=False, float_format=lambda value: f"{value:.3f}")

    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>NHANES 心代谢风险数据挖掘报告</title>
  <style>
    body {{ margin: 32px; font-family: Arial, 'Microsoft YaHei', sans-serif; color: #1f2937; line-height: 1.6; }}
    h1, h2 {{ color: #0f172a; }}
    .grid {{ display: grid; grid-template-columns: repeat(4, minmax(0, 1fr)); gap: 12px; margin: 20px 0; }}
    .metric {{ border: 1px solid #d1d5db; border-radius: 6px; padding: 12px; background: #f9fafb; }}
    .metric strong {{ display: block; font-size: 22px; color: #111827; }}
    table {{ border-collapse: collapse; width: 100%; font-size: 13px; margin-bottom: 20px; }}
    th, td {{ border: 1px solid #d1d5db; padding: 6px 8px; text-align: left; }}
    th {{ background: #e5e7eb; }}
    img {{ max-width: 100%; border: 1px solid #e5e7eb; border-radius: 6px; margin: 10px 0 24px; }}
  </style>
</head>
<body>
  <h1>NHANES 心代谢风险分层、关联规则与分类预测报告</h1>
  <p>本报告由项目流水线自动生成，覆盖数据预处理、数据仓库、聚合分析、Apriori 关联规则和决策树分类。</p>
  <div class="grid">
    <div class="metric">成人样本<strong>{len(processed):,}</strong></div>
    <div class="metric">高风险比例<strong>{processed['high_cardiometabolic_risk'].mean():.1%}</strong></div>
    <div class="metric">分类准确率<strong>{metrics.get('accuracy', 0):.3f}</strong></div>
    <div class="metric">分类 F1<strong>{metrics.get('f1', 0):.3f}</strong></div>
  </div>
  <h2>聚合分析</h2>
  <h3>按年龄分组</h3>
  {age_html}
  <img src="aggregation_risk_by_age.png" alt="Risk by age">
  <img src="aggregation_disease_rates_by_age.png" alt="Disease rates by age">
  <h3>按 BMI 分组</h3>
  {bmi_html}
  <img src="aggregation_risk_by_bmi.png" alt="Risk by BMI">
  <h2>关联规则</h2>
  {top_rules_html}
  <h2>分类结果</h2>
  {metrics_html}
  <img src="classification_confusion_matrix.png" alt="Confusion matrix">
  <h3>特征重要性</h3>
  {imports_html}
  <img src="classification_feature_importance.png" alt="Feature importance">
</body>
</html>"""
    report_path.write_text(html, encoding="utf-8")
    return report_path


def write_tree_text(tree_text: str, output_dir: Path) -> Path:
    path = output_dir / "classification_tree_rules.txt"
    path.write_text(tree_text, encoding="utf-8")
    return path


def run_pipeline(
    data_path: Optional[Path] = None,
    output_dir: Path = RESULTS_DIR,
    min_support: float = 0.08,
    min_confidence: float = 0.55,
    max_itemset_len: int = 3,
    max_depth: int = 5,
    nrows: Optional[int] = None,
) -> PipelineArtifacts:
    ensure_dirs()
    output_dir.mkdir(parents=True, exist_ok=True)

    raw_df = load_raw_data(data_path=data_path, nrows=nrows)
    raw_subject = aggregate_subject_level(raw_df)
    processed = add_health_features(raw_subject, adult_only=True)
    aggregations = build_aggregation_tables(processed)
    transactions = [transaction_from_row(row) for _, row in processed.iterrows()]
    itemsets, rules = apriori_rules(
        transactions,
        min_support=min_support,
        min_confidence=min_confidence,
        max_len=max_itemset_len,
    )
    metrics, confusion, importances, tree_text = run_classification(processed, max_depth=max_depth)

    save_csv_outputs(processed, aggregations, itemsets, rules, metrics, confusion, importances, output_dir)
    save_visualizations(aggregations, confusion, importances, output_dir)
    warehouse_path = write_sqlite_warehouse(
        raw_subject,
        processed,
        aggregations,
        itemsets,
        rules,
        metrics,
        confusion,
        importances,
        output_dir,
    )
    generate_warehouse_design(processed, output_dir)
    report_md_path, report_docx_path = generate_course_report(
        processed,
        aggregations,
        rules,
        metrics,
        importances,
        output_dir,
    )
    report_pdf_path = generate_pdf_report(
        processed,
        aggregations,
        rules,
        metrics,
        importances,
        output_dir,
    )
    report_html_path = generate_html_report(processed, aggregations, rules, metrics, importances, output_dir)
    write_tree_text(tree_text, output_dir)

    summary = {
        "raw_rows": int(len(raw_df)),
        "subject_rows": int(len(raw_subject)),
        "processed_rows": int(len(processed)),
        "association_rules": int(len(rules)),
        "high_risk_rate": float(processed["high_cardiometabolic_risk"].mean()),
        "classification": metrics,
        "output_dir": str(output_dir),
    }
    (output_dir / "pipeline_summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    return PipelineArtifacts(
        output_dir=output_dir,
        warehouse_path=warehouse_path,
        processed_path=output_dir / "nhanes_health_subjects.csv",
        rules_path=output_dir / "association_rules.csv",
        classification_metrics_path=output_dir / "classification_metrics.csv",
        report_html_path=report_html_path,
        report_md_path=report_md_path,
        report_docx_path=report_docx_path,
        report_pdf_path=report_pdf_path,
        summary=summary,
    )


def main(argv: Optional[Sequence[str]] = None) -> PipelineArtifacts:
    parser = argparse.ArgumentParser(description="Run the NHANES medical data mining pipeline.")
    parser.add_argument("--data", type=Path, default=None, help="CSV dataset path. Defaults to data/nhanes_processed.csv.")
    parser.add_argument("--min-support", type=float, default=0.08, help="Minimum support for Apriori.")
    parser.add_argument("--min-confidence", type=float, default=0.55, help="Minimum confidence for Apriori rules.")
    parser.add_argument("--max-depth", type=int, default=5, help="Decision tree max depth.")
    parser.add_argument("--nrows", type=int, default=None, help="Optional row limit for quick testing.")
    args = parser.parse_args(argv)

    artifacts = run_pipeline(
        data_path=args.data,
        min_support=args.min_support,
        min_confidence=args.min_confidence,
        max_depth=args.max_depth,
        nrows=args.nrows,
    )
    print(json.dumps(artifacts.summary, ensure_ascii=False, indent=2))
    print(f"Warehouse: {artifacts.warehouse_path}")
    print(f"HTML report: {artifacts.report_html_path}")
    print(f"Markdown report: {artifacts.report_md_path}")
    if artifacts.report_docx_path:
        print(f"Word report: {artifacts.report_docx_path}")
    if artifacts.report_pdf_path:
        print(f"PDF report: {artifacts.report_pdf_path}")
    return artifacts


if __name__ == "__main__":
    main()
